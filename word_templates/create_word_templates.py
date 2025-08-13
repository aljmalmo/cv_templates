#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
سكريبت لإنشاء قوالب Word للسير الذاتية
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_cv_template(field_name, field_color, sample_data):
    """إنشاء قالب Word للسيرة الذاتية"""
    
    # إنشاء مستند جديد
    doc = Document()
    
    # تعديل الهوامش
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # إضافة عنوان رئيسي
    title = doc.add_heading(sample_data['name'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(*field_color)
    
    # إضافة المسمى الوظيفي
    subtitle = doc.add_paragraph(sample_data['title'])
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # إضافة معلومات الاتصال
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f"{sample_data['email']} | {sample_data['phone']} | {sample_data['location']}"
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(12)
    
    # إضافة خط فاصل
    doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # إضافة النبذة المختصرة
    summary_heading = doc.add_heading('نبذة مختصرة', level=1)
    summary_heading.runs[0].font.color.rgb = RGBColor(*field_color)
    doc.add_paragraph(sample_data['summary'])
    
    # إضافة الخبرة المهنية
    exp_heading = doc.add_heading('الخبرة المهنية', level=1)
    exp_heading.runs[0].font.color.rgb = RGBColor(*field_color)
    
    for exp in sample_data['experience']:
        job_para = doc.add_paragraph()
        job_run = job_para.add_run(exp['title'])
        job_run.bold = True
        job_run.font.size = Pt(14)
        
        company_para = doc.add_paragraph(f"{exp['company']} | {exp['date']}")
        company_para.runs[0].font.color.rgb = RGBColor(*field_color)
        
        doc.add_paragraph(exp['description'])
        doc.add_paragraph()  # مسافة فارغة
    
    # إضافة التعليم
    edu_heading = doc.add_heading('التعليم', level=1)
    edu_heading.runs[0].font.color.rgb = RGBColor(*field_color)
    
    for edu in sample_data['education']:
        degree_para = doc.add_paragraph()
        degree_run = degree_para.add_run(edu['degree'])
        degree_run.bold = True
        degree_run.font.size = Pt(14)
        
        institution_para = doc.add_paragraph(f"{edu['institution']} | {edu['date']}")
        institution_para.runs[0].font.color.rgb = RGBColor(*field_color)
        
        if 'description' in edu:
            doc.add_paragraph(edu['description'])
        doc.add_paragraph()  # مسافة فارغة
    
    # إضافة المهارات
    skills_heading = doc.add_heading('المهارات', level=1)
    skills_heading.runs[0].font.color.rgb = RGBColor(*field_color)
    
    for skill_category in sample_data['skills']:
        category_para = doc.add_paragraph()
        category_run = category_para.add_run(skill_category['category'])
        category_run.bold = True
        category_run.font.size = Pt(12)
        
        skills_text = " • ".join(skill_category['items'])
        doc.add_paragraph(skills_text)
    
    # إضافة اللغات
    lang_heading = doc.add_heading('اللغات', level=1)
    lang_heading.runs[0].font.color.rgb = RGBColor(*field_color)
    
    for lang in sample_data['languages']:
        lang_para = doc.add_paragraph(f"{lang['name']}: {lang['level']}")
    
    # إضافة الشهادات
    cert_heading = doc.add_heading('الشهادات والدورات', level=1)
    cert_heading.runs[0].font.color.rgb = RGBColor(*field_color)
    
    for cert in sample_data['certifications']:
        cert_para = doc.add_paragraph()
        cert_run = cert_para.add_run(cert['name'])
        cert_run.bold = True
        
        issuer_para = doc.add_paragraph(f"{cert['issuer']} | {cert['date']}")
        issuer_para.runs[0].font.color.rgb = RGBColor(*field_color)
    
    return doc

# بيانات نموذجية لكل مجال
templates_data = {
    'tech': {
        'name': 'أحمد محمد العلي',
        'title': 'مطور برمجيات أول',
        'email': 'ahmed.ali@email.com',
        'phone': '+966 50 123 4567',
        'location': 'الرياض، المملكة العربية السعودية',
        'summary': 'مطور برمجيات متخصص مع أكثر من 5 سنوات من الخبرة في تطوير تطبيقات الويب والهاتف المحمول. خبرة واسعة في JavaScript، React، Node.js، وقواعد البيانات.',
        'experience': [
            {
                'title': 'مطور برمجيات أول',
                'company': 'شركة التقنية المتقدمة',
                'date': 'مارس 2021 - حتى الآن',
                'description': 'قيادة فريق من 4 مطورين في تطوير تطبيقات ويب متقدمة باستخدام React و Node.js.'
            }
        ],
        'education': [
            {
                'degree': 'بكالوريوس علوم الحاسب',
                'institution': 'جامعة الملك سعود',
                'date': '2019',
                'description': 'المعدل التراكمي: 3.8/4.0'
            }
        ],
        'skills': [
            {'category': 'لغات البرمجة', 'items': ['JavaScript', 'TypeScript', 'Python', 'Java']},
            {'category': 'تقنيات الواجهة الأمامية', 'items': ['React', 'Vue.js', 'HTML5', 'CSS3']}
        ],
        'languages': [
            {'name': 'العربية', 'level': 'لغة أم'},
            {'name': 'الإنجليزية', 'level': 'طليق'}
        ],
        'certifications': [
            {'name': 'AWS Certified Developer', 'issuer': 'Amazon Web Services', 'date': 'مايو 2023'}
        ]
    },
    
    'marketing': {
        'name': 'فاطمة أحمد الزهراني',
        'title': 'أخصائي التسويق الرقمي',
        'email': 'fatima.zahrani@email.com',
        'phone': '+966 55 987 6543',
        'location': 'جدة، المملكة العربية السعودية',
        'summary': 'أخصائي تسويق رقمي مبدع مع خبرة 4 سنوات في إدارة الحملات التسويقية وزيادة المبيعات.',
        'experience': [
            {
                'title': 'أخصائي التسويق الرقمي',
                'company': 'شركة الإبداع التسويقي',
                'date': 'يناير 2022 - حتى الآن',
                'description': 'إدارة الحملات التسويقية الرقمية عبر منصات متعددة.'
            }
        ],
        'education': [
            {
                'degree': 'بكالوريوس إدارة الأعمال - تخصص تسويق',
                'institution': 'جامعة الملك عبدالعزيز',
                'date': '2020'
            }
        ],
        'skills': [
            {'category': 'التسويق الرقمي', 'items': ['Google Ads', 'Facebook Ads', 'Instagram Marketing']},
            {'category': 'تحليل البيانات', 'items': ['Google Analytics', 'Facebook Insights']}
        ],
        'languages': [
            {'name': 'العربية', 'level': 'لغة أم'},
            {'name': 'الإنجليزية', 'level': 'متقدم'}
        ],
        'certifications': [
            {'name': 'Google Ads Certified', 'issuer': 'Google', 'date': 'مارس 2023'}
        ]
    }
}

# ألوان المجالات
field_colors = {
    'tech': (44, 62, 80),      # أزرق داكن
    'marketing': (39, 174, 96), # أخضر
    'design': (142, 68, 173),   # بنفسجي
    'medical': (192, 57, 43),   # أحمر
    'engineering': (230, 126, 34), # برتقالي
    'finance': (52, 73, 94),    # رمادي داكن
    'education': (22, 160, 133), # تركوازي
    'legal': (44, 62, 80)       # أسود
}

def main():
    """الدالة الرئيسية لإنشاء جميع القوالب"""
    
    # إنشاء مجلد الإخراج
    output_dir = '/home/ubuntu/cv_templates_github/word_templates'
    os.makedirs(output_dir, exist_ok=True)
    
    # إنشاء قوالب تكنولوجيا المعلومات والتسويق كنماذج
    for field, data in templates_data.items():
        doc = create_cv_template(field, field_colors[field], data)
        filename = f'{field}_template.docx'
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        print(f"تم إنشاء قالب {field}: {filename}")

if __name__ == "__main__":
    main()

